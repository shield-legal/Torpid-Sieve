from docx import Document

def parse_word_document(file_path):
    # Load the document
    doc = Document(file_path)
    
    # Initialize a dictionary to store text by section
    parsed_content = {"paragraphs": [], "tables": []}
    
    # Extract text from paragraphs
    for para in doc.paragraphs:
        parsed_content["paragraphs"].append(para.text)
    
    # Extract text from tables
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        parsed_content["tables"].append(table_data)
    
    return parsed_content

# Usage
file_path = "C:\\Users\\Ahsan\\Desktop\\Work Code\\XML Parser\\Plaintiff Intakes\\Plaintiff Intake Alicia Belli AB-137.docx"
content = parse_word_document(file_path)

# Display content
print("Paragraphs:")
for paragraph in content["paragraphs"]:
    print(paragraph)

print("\nTables:")
for table in content["tables"]:
    for row in table:
        print(row)
